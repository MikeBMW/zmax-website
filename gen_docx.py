#!/usr/bin/env python3
"""Z700 立项申请书 DOCX — 从 DDS 数据库读取，指定中文字体"""
import sqlite3, datetime, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DB = '/www/wwwroot/datadrive.world/dds.db'

def set_cn_font(run, name='等线', size=11, bold=False, color=None):
    """设置中文字体 — 关键：必须同时设西文和东亚字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    # 设置东亚字体（中文）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if color:
        run.font.color.rgb = color

conn = sqlite3.connect(DB)

def q(table, key, col='value'):
    r = conn.execute(f"SELECT {col} FROM {table} WHERE key=?", (key,)).fetchone()
    return r[0] if r else ''

def qrow(table, key_col='id'):
    c = conn.cursor()
    c.execute(f"SELECT * FROM {table}")
    cols = [d[0] for d in c.description]
    result = {}
    for r in c.fetchall():
        d = dict(zip(cols, r))
        k = d.pop(key_col)
        result[k] = d
    return result

doc = Document()

# Set default font via style
style = doc.styles['Normal']
style.font.name = '等线'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
# Set east-asia font on style
rPr = style.element.get_or_add_rPr()
rFonts = style.element.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), '等线')
rFonts.set(qn('w:ascii'), '等线')
rFonts.set(qn('w:hAnsi'), '等线')
rPr.insert(0, rFonts)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_cn_font(run, '等线', 16 if level==1 else 13, bold=True, color=RGBColor(0x00,0x00,0x00))
    return h

def para(text, bold=False, sz=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, '等线', sz, bold, color)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    # Clear default run and re-add
    p.clear()
    run = p.add_run(text)
    set_cn_font(run, '等线', 10.5)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, '等线', 10, bold=True, color=RGBColor(0x00,0x00,0x00))
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_cn_font(run, '等线', 10)
    doc.add_paragraph()
    return t

# ═══════════ COVER ═══════════
for _ in range(5):
    doc.add_paragraph()
para("Z700 光模块精密制造", bold=True, sz=26, color=RGBColor(0x00,0x00,0x00))
para("具身智能机器人系统 · 立项申请书", bold=True, sz=18)
doc.add_paragraph()
para(f"申请单位：智蜂创元 ZFCY")
para("技术路线：视触觉混合动作模型 · 端到端感知决策控制")
para("应用场景：光模块（100G/400G/800G）精密制造产线")
para(f"申请日期：{datetime.datetime.now().strftime('%Y年%m月%d日')}")
para("文档编号：ZFCY-Z700-001  |  版本：V1.0  |  密级：内部")
doc.add_page_break()

# ═══════════ 1. 项目概述 ═══════════
heading("一、项目概述")

heading("1.1 项目背景与市场机会", 2)
para("光模块是AI算力基础设施的核心器件，2026年全球市场规模超200亿美元。随着800G/1.6T高速光模块进入规模量产阶段，精密制造能力成为产能瓶颈。传统自动化设备无法应对换型频繁、精度要求极高的光模块制造，具身智能技术成熟度已达到产业化临界点，具备替代人工精细操作的能力。")

heading("1.2 项目目标", 2)
para("开发面向光模块工厂的具身智能机器人系统 Z700，实现 ±0.02mm 精密操作，关键工序良率 >99%。覆盖 COC/OE/MOD/WH 四区全流程。从 L2 单工位到 L4 全产线的软件定义进化之路。")

heading("1.3 核心创新", 2)
for item in ["视触觉混合动作模型：多模态感知→决策→控制的端到端架构",
             ">1kHz 毫秒级力控闭环：金手指零损伤",
             "零样本泛化：换型无需重新训练，即插即用",
             "数据闭环：每次操作都是训练数据，产线越用越聪明"]:
    bullet(item)
doc.add_page_break()

# ═══════════ 2. 产品架构 ═══════════
heading("二、产品与系统架构")

heading("2.1 Z700 轮式双臂机器人 [L4旗舰]", 2)
para("自主导航底盘 + 六轴双臂协同。覆盖老化箱插拔、工位间物料流转、产线巡检。全自主执行+自适应换型+自恢复异常。目标：2027Q2 全产线智能协同。")

heading("2.2 Z700F 固定式精密插拔 [L2基线]", 2)
para("SR5-C六轴机械臂 + DH夹爪。5工位循环：入料→扫码→刷程序→AOI检测→出料。>1kHz毫秒级力控闭环。当前状态：已部署验证，99%+良率。")

heading("2.3 系统架构", 2)
table(["层级", "功能", "核心技术"], [
    ["感知层", "多模态信号融合", "RealSense D405 + 腕部RGB + 六维力/力矩 + 触觉阵列"],
    ["决策层", "VLA推理", "视触觉混合动作模型，端到端像素到动作，<15ms推理"],
    ["执行层", "力控闭环", ">1kHz力控 + Lissajous力搜索，±0.02mm精密定位"],
    ["进化层", "数据闭环", "端侧采集→云端训练→模型部署，持续迭代优化"],
])
doc.add_page_break()

# ═══════════ 3. 核心技术 ═══════════
heading("三、核心技术能力")
table(["模块", "技术要点", "核心指标"], [
    ["场景感知", "D405深度+腕部RGB+力/触觉，四路传感器实时融合", "毫米级精度"],
    ["智能决策", "视触觉混合动作模型，端到端像素到动作", "<15ms推理"],
    ["运动控制", "关节+末端+指尖三路力反馈，Lissajous力搜索", ">1kHz闭环"],
    ["精细操作", "双臂协同+触觉反馈，全系列光模块适配", "金手指零损伤"],
    ["数据闭环", "端侧采集→云端训练→模型部署，7x24流水线", "持续迭代"],
    ["泛化稳定", "世界模型预测物理交互，五层安全保护", ">99%稳定"],
])
doc.add_page_break()

# ═══════════ 4. 工厂布局 ═══════════
heading("四、目标工厂：800G DR8 光模块产线")

heading("4.1 四区布局", 2)
table(["区域", "工位", "编号", "核心工序", "机器人"],
    [["COC基板区","24站","I101-I124","COC贴片、金线键合","Z700轮式双臂巡检"],
     ["OE光引擎区","72站","I125-I196","DA贴片、Plasma、烘烤、烧录、AOI","Z100L料笼搬运"],
     ["MOD模块区","32站","I197-I228","模块耦合、老化、眼图测试","Z700F+AOI精密目检"],
     ["WH仓库区","N站","I223-I228+I500","物料存储、AGV调度、出厂终检","Z100L上下料搬运"]])

heading("4.2 机器人部署", 2)
para("子工序级精确分配：目检/镜检 → Z700F+AOI；上下料/料笼搬运 → Z100L；自动设备 → 无需机器人。每台机器人覆盖5-8个相邻工位。OE区统一配Z100L，理由+操作对象全程可追溯。")
doc.add_page_break()

# ═══════════ 5. 合作模式 ═══════════
heading("五、工程合作模式：智蜂×它石/他山")

heading("5.1 智蜂 (ZFCY) — AI模型与智能引擎", 2)
for item in ["视触觉混合动作模型研发","感知/决策/控制算法迭代","数据闭环与模型训练","仿真环境与数字孪生","云端训练基础设施","技术路线制定与攻关"]:
    bullet(item)

heading("5.2 它石/他山 — 硬件集成与工程落地", 2)
for item in ["机械臂/夹爪/传感器选型集成","工位改造与产线适配","电气/气动/通信系统集成","现场部署调试与验收","售后运维与技术支持","生产制造与供应链管理"]:
    bullet(item)

heading("5.3 ASPICE V-Model 协同开发", 2)
para("左翼（需求→设计）：智蜂主导系统需求、架构设计、AI模型规格。谷底（实现）：智蜂负责模型训练与算法开发，它石负责硬件集成与工装制造。右翼（验证→验收）：联合单元测试、集成测试、系统验证、现场验收。RASI责任矩阵 + WBS分解 + 三方供应能力对比，确保供应链安全。")
doc.add_page_break()

# ═══════════ 6. 开发阶段 ═══════════
heading("六、开发阶段与里程碑")
table(["阶段", "时间", "周期", "关键交付", "验收标准"], [
    ["Phase 1\n单工位验证","2026 Q3","3个月",
     "Z700F固定式插拔产线\nSR5-C机械臂集成\n5工位循环自动化\n力控闭环算法固件",
     "力控±0.02mm\n良率>99%"],
    ["Phase 2\n多工位联线","2026 Q4","3个月",
     "Z700轮式双臂机器人\n自主导航+双臂协同\n多工位物料流转\nVLA模型换型泛化",
     "3-5工位联线\n换型<5分钟"],
    ["Phase 3\n全产线智能","2027 Q1-Q2","6个月",
     "全产线机器人部署\n多机协同调度\n端到端数据闭环\nOTA升级系统",
     "7x24无人值守\n良率>99%\n换型零干预"],
])
doc.add_page_break()

# ═══════════ 7. 交付物 ═══════════
heading("七、交付物清单")
table(["阶段","类别","交付内容"], [
    ["Phase 1","硬件","Z700F固定式插拔产线x1套；SR5-C机械臂集成方案；DH夹爪/力传感器/视觉模组；工位改造与工装设计"],
    ["Phase 1","软件","力控闭环算法固件；5工位循环控制程序；AOI检测模型；边缘推理部署包"],
    ["Phase 2","硬件","Z700轮式双臂机器人x1台；自主导航底盘；双臂协同控制箱；Z100L料笼搬运AGV"],
    ["Phase 2","软件","VLA混合动作模型；自主导航SLAM系统；多工位调度中间件；换型泛化配置工具"],
    ["Phase 3","硬件","全产线机器人联调方案；多机协同通信网络；安全防护系统；备品备件清单"],
    ["Phase 3","软件","全产线调度系统；数据闭环训练平台；OTA升级管理系统；运维监控仪表盘"],
    ["全阶段","文档","系统设计说明书；接口协议文档；操作维护手册；ASPICE V-Model全套；验收测试报告"],
    ["全阶段","培训","操作员培训(2天)；维护工程师培训(3天)；产线管理者培训(1天)；远程技术支持(12个月)"],
])
doc.add_page_break()

# ═══════════ 8. 投资回报 ═══════════
heading("八、投资预估与回报分析")

heading("8.1 投资预估（12个月）", 2)
table(["项目","Phase1(Q3)","Phase2(Q4)","Phase3(Q1-Q2)","合计"], [
    ["硬件投入","待确认","待确认","待确认","待确认"],
    ["软件研发","待确认","待确认","待确认","待确认"],
    ["云GPU算力","待确认","待确认","待确认","待确认"],
    ["人员成本","待确认","待确认","待确认","待确认"],
])

heading("8.2 回报分析", 2)
para("直接效益：单工位替代2-3名操作员 × 72工位 = 显著人力节省；良率从人工~95%提升至>99%，减少返工损耗；7x24连续运行 vs 人工2班倒，产能翻倍。")
para("间接效益：换型时间从数天缩短到分钟级；数据驱动工艺优化，持续提升良率；减少对熟练工人的依赖，降低培训成本；建立行业技术壁垒。")
para("预计投资回收期：待确认个月")

heading("8.3 风险与对策", 2)
table(["风险","对策"], [
    ["技术风险：VLA模型真实产线泛化能力","数字孪生+仿真训练+渐进式部署"],
    ["工程风险：工位改造与产线兼容性","模块化工装+最小化停线时间"],
    ["市场风险：光模块技术路线迭代","模型零样本泛化，换型无需重新训练"],
    ["供应链风险：核心传感器依赖进口","它石/他山双供应商+国产替代"],
])
doc.add_page_break()

# ═══════════ 9. 团队 ═══════════
heading("九、核心团队与技术储备")

heading("9.1 核心团队", 2)
table(["角色","成员","职责"], [
    ["技术负责人/总工","xspace","系统架构总设计；VLA模型/ACT/LeWM；多引擎协同推理；技术战略规划"],
    ["产品经理/模型专家","web","VLA大模型训练验证；GR00T/仿真/MuJoCo；ComfyUI数据管线；产品规划与项目管理"],
    ["硬件/安全/测试","小芳","Orin Nano端侧部署；Mac中转/硬件集成；ROS2/传感器采集；安全测试与验收"],
])

heading("9.2 技术储备", 2)
para("模型引擎：五引擎切换（SmolVLA / GR00T-N1.7 / ACT / LeWM / VLA-T）")
para("训练设施：RTX 4090 24GB云端训练 · RTX 4060 8GB端侧推理")
para("硬件平台：Orin Nano + RealSense D405 + SR5-C六轴机械臂 + DH夹爪 + 六维力传感器")
para("软件工具链：MuJoCo仿真 · RoboGen数据生成 · ComfyUI/DDS数据空间 · MCAP数据采集")
para("开发体系：ASPICE V-Model · RASI矩阵 · WBS分解 · Git三工程协同")
para("数据闭环：Orin端侧采集→MAC中转→ECS→4090训练→模型部署回端侧，全链路打通")
doc.add_page_break()

# ═══════════ 10. 总结 ═══════════
heading("十、立项总结")

heading("10.1 核心价值", 2)
para("技术先进性：视触觉混合动作模型、端到端感知决策控制、>1kHz力控闭环、零样本泛化。")
para("市场紧迫性：AI算力爆发驱动光模块扩产，精密制造人工替代刚需，竞争对手尚在实验室阶段。")
para("团队完备性：三人核心团队覆盖全栈，模型+硬件+部署闭环，真实产线数据已验证。")

heading("10.2 下一步行动", 2)
table(["时间","行动"], [
    ["本月","完成立项审批；确定合作方（它石/他山）；签署NDA与技术协议"],
    ["Q3 (Phase1)","Z700F固定式产线验收；力控精度±0.02mm达标；良率>99%验收确认"],
    ["Q4 (Phase2)","Z700轮式双臂部署；多工位联线验证；VLA模型换型泛化测试"],
    ["2027 Q1-Q2 (Phase3)","全产线多机协同部署；数据闭环完整运行；7x24无人值守验收交付"],
])

doc.add_paragraph()
para("datadrive.world  |  智蜂创元 ZFCY  |  光模块精密制造 · 具身智能", sz=9)
conn.close()

out = '/www/wwwroot/datadrive.world/Z700-立项申请书.docx'
doc.save(out)
print(f"Saved: {out} ({os.path.getsize(out)/1024:.1f} KB)")
